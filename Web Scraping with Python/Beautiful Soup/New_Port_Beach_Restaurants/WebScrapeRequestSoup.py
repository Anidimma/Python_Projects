
#IMPORTING THE MODULES
import requests
from bs4 import *
from docx import *
import pandas as pd 


columns = ['Name','Rating', 'Number','Location']
hold = []
       
general_doc = Document()
images_doc = Document()

# SPECIFYING HOW THE CODE SHOULD ACCESS THE WEBSITE
"""To understand how the variables under work with the code please
go to 'www.yelp.com' search for 'Resturants' and for location put 'Newport Beach, CA'
after that you will be taken to a new webpage, observe the url of that
web page then you'll understand how the codes below work with the defined variables"""

base_url = 'https://www.yelp.com/search?find_desc=Restaurants&find_loc=' # Initial url
loc = 'Newport Beach, CA' # Specifying the city
sorter = 'recommended'
current_page = 0

"""The arrangement of the various code below is to beautify the arrangement on microsoft word"""

all_values = general_doc.add_heading('THIS IS THE DATA SCRAPPED FROM YELP WEBSITE')
all_values = general_doc.add_paragraph("""""")

# GATHERING DATA FROM 200 SOURCES APART FROM IMAGE DATA   
while current_page < 201:

    url = base_url + loc + '&sortby='+ sorter + '&start=' + str(current_page)
    yelp_r = requests.get(url)
    yelp_soup = BeautifulSoup(yelp_r.text, 'html.parser')

    #SPECIFYING WHERE ON THE PAGE PYTHON SHOULD EXTRACT DATA FROM
    general = yelp_soup.findAll(class_ ='arrange-unit__09f24__1gZC1 arrange-unit-fill__09f24__O6JFU border-color--default__09f24__R1nRO')

    print('\n',f'You\'re on page {current_page}','\n')

    for num,biz in enumerate(general):
        if len(biz) > 1:
            for num2,val in enumerate(biz):
                if num2 == 0:
                    val = str(val)
                    holder = val.split('">(')
                    if len(holder) == 2:
                        
                        val1 = val.split('name=')[1].split('"')[1].replace('&amp;','&')  # Getting the name
                        print(val1)
                        all_values.add_run(f""" 

                        This is from page {current_page}  
                        --------------------------------------------------------------------------------------------------------------

                        {val1}""")

                        val2 = val.split('aria-label=')[1].split('"')[1]  # Getting the rating
                        print(val2)
                        all_values.add_run(f""" 

                        {val2}""")

                        val3 = val.split('">(')[1].split('</p></div>')[0]  # Getting the phone number
                        val3 = '(' + val3 
                        print(val3)
                        all_values.add_run(f""" 

                        {val3}""")

                        val4 = val.split('buy">')[1].split('</span></p>')[0]  # Getting the location
                        print(val4)
                        all_values.add_run(f""" 

                        {val4}

                        """)

                        container = [val1,val2,val3,val4]
                        hold.append(container)
                        print('\n')
                    
                    elif len(holder) == 1:
                        val1 = val.split('name=')[1].split('"')[1].replace('&amp;','&')  # Getting the name
                        print(val1)
                        all_values.add_run(f""" 

                        This is from page {current_page}  
                        --------------------------------------------------------------------------------------------------------------

                        {val1}""")

                        val2 = val.split('aria-label=')[1].split('"')[1]  # Getting the rating
                        print(val2)
                        all_values.add_run(f""" 

                        {val2}""")

                        val3 = 'N/A'  # When no phone number available
                        print(val3)
                        all_values.add_run(f""" 

                        {val3}""")
                        
                        val4 = val.split('buy">')[1].split('</span></p>')[0]  # Getting the location
                        print(val4)
                        all_values.add_run(f""" 

                        {val4}

                        """)

                        container = [val1,val2,val3,val4]
                        hold.append(container)
                        print('\n')

    general_doc.save(r'C:/Users/Divinity-PC/Desktop/practice/np_resturants.docx')  # Specifying the place where the data will be saved             
    print('\n',f'You left page {current_page}','\n')

    current_page += 10
general_doc = Document() # Resetting the document place holder
print('stop')

df = pd.DataFrame(hold,columns = columns)
df = df.drop_duplicates(subset= 'Name', keep="first", inplace=False, ignore_index = True) #Dropping Duplicate Entries 
df.to_csv(r'C:/Users/Divinity-PC/Desktop/practice/np_resturants.csv', index=False)

